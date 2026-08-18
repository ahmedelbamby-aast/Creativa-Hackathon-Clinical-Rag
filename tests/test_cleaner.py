from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from chunking.src.services.ingestion import cleaner


@pytest.mark.parametrize(
    ("name", "content_type", "extractor"),
    [
        ("report.PDF", "application/pdf", "pypdf"),
        ("report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "python-docx"),
        ("notes.txt", "text/plain", "plain-text"),
        ("unknown.bin", "text/plain", "plain-text"),
    ],
)
def test_content_type_and_extractor_detection(name, content_type, extractor):
    assert cleaner.detect_content_type(name) == content_type
    assert cleaner.get_extractor_label(name) == extractor


def test_unknown_and_extensionless_files_share_plain_text_fallback(tmp_path):
    for path in (tmp_path / "unknown.bin", tmp_path / "README"):
        assert cleaner.detect_content_type(path) == "text/plain"
        assert cleaner.get_extractor_label(path) == "plain-text"


def test_extract_utf8_and_latin1_text(tmp_path):
    utf8 = tmp_path / "utf8.txt"
    utf8.write_text("  diabetes café  ", encoding="utf-8")
    latin1 = tmp_path / "latin1.txt"
    latin1.write_bytes("café".encode("latin-1"))
    assert cleaner.extract_pages(str(utf8)) == [{"page_number": 1, "text": "diabetes café"}]
    assert cleaner.extract_pages(str(latin1)) == [{"page_number": 1, "text": "café"}]


def test_extract_docx(tmp_path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Diabetes care")
    document.add_paragraph("   ")
    document.add_paragraph("Second paragraph")
    document.save(path)
    assert cleaner.extract_pages(str(path)) == [
        {"page_number": 1, "text": "Diabetes care\nSecond paragraph"}
    ]


def test_pdf_extractor_handles_textless_pages_and_reader_failures(monkeypatch):
    class Page:
        def __init__(self, text):
            self.text = text

        def extract_text(self):
            return self.text

    class Reader:
        def __init__(self, _):
            self.pages = [Page("first"), Page(None)]

    monkeypatch.setattr("pypdf.PdfReader", Reader)
    assert cleaner.extract_pages("sample.pdf") == [
        {"page_number": 1, "text": "first"},
        {"page_number": 2, "text": ""},
    ]

    monkeypatch.setattr("pypdf.PdfReader", lambda _: (_ for _ in ()).throw(OSError("bad pdf")))
    assert cleaner.extract_pages("bad.pdf") == []


def test_docx_extractor_failure_returns_no_pages(tmp_path):
    assert cleaner.extract_pages(str(tmp_path / "missing.docx")) == []


def test_extract_missing_or_empty_file_returns_no_pages(tmp_path):
    assert cleaner.extract_pages(str(tmp_path / "missing.txt")) == []
    empty = tmp_path / "empty.txt"
    empty.write_text("  ", encoding="utf-8")
    assert cleaner.extract_pages(str(empty)) == []


def test_extract_pages_rejects_unknown_content_type(monkeypatch):
    monkeypatch.setattr(cleaner, "detect_content_type", lambda _: "application/unknown")
    assert cleaner.extract_pages("unknown.data") == []


def test_clean_pages_removes_artifacts_and_tracks_statistics():
    pages = [
        {"page_number": index, "text": f"\ufeffCOMMON HEADER\npre-\nvention\u00a0text {index}\n{index}\nsite.example | {index}\n\n\n"}
        for index in range(1, 5)
    ]
    cleaned, stats = cleaner.clean_pages(pages)
    assert [page["page_number"] for page in cleaned] == [1, 2, 3, 4]
    assert [page["text"] for page in cleaned] == [f"prevention text {index}" for index in range(1, 5)]
    assert stats["page_number_lines_removed"] == 4
    assert stats["repeated_boilerplate_lines_removed"] == 4
    assert stats["footer_lines_removed"] == 4
    assert stats["chars_after"] < stats["chars_before"]


def test_extract_text_joins_cleaned_pages(monkeypatch):
    monkeypatch.setattr(cleaner, "extract_pages", lambda _: [
        {"page_number": 1, "text": "first"},
        {"page_number": 2, "text": "second"},
    ])
    assert cleaner.extract_text("ignored.pdf") == "first\n\nsecond"


def test_repeated_line_cleaner_handles_small_and_nonrepeating_inputs():
    small = ["header\nbody", "header\nother"]
    assert cleaner._remove_repeated_lines(small) == (small, 0, 0)
    distinct = [f"unique content {index}" for index in range(4)]
    assert cleaner._remove_repeated_lines(distinct) == (distinct, 0, 0)


def test_repeated_footer_is_removed_when_glued_to_content():
    footer = "long-running-footer.example"
    pages = [
        f"{footer} | {index}\ncontent {index} {footer}\nkeep {index}"
        for index in range(1, 5)
    ]
    cleaned, exact_removed, marker_removed = cleaner._remove_repeated_lines(pages)
    assert cleaned == [f"content {index}\nkeep {index}" for index in range(1, 5)]
    assert exact_removed == 0
    assert marker_removed == 8


def test_clean_pages_discards_empty_pages():
    cleaned, stats = cleaner.clean_pages([{"page_number": 1, "text": "  \n  "}])
    assert cleaned == []
    assert stats["pages_after"] == 0


def test_extract_text_returns_none_for_extraction_or_cleaning_failure(monkeypatch):
    monkeypatch.setattr(cleaner, "extract_pages", lambda _: [])
    assert cleaner.extract_text("missing.txt") is None
    monkeypatch.setattr(cleaner, "extract_pages", lambda _: [{"page_number": 1, "text": "text"}])
    monkeypatch.setattr(cleaner, "clean_pages", lambda _: ([], {}))
    assert cleaner.extract_text("empty.txt") is None
